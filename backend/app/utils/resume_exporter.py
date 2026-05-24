import io
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def generate_resume_pdf(resume_data: dict) -> io.BytesIO:
    buffer = io.BytesIO()
    
    # 0.5 inch margins (36 pt) for standard clean layout
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    
    styles = getSampleStyleSheet()
    
    # Custom Typography styles for polished look
    title_style = ParagraphStyle(
        'ResumeTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#111827'),
        alignment=1, # Center
        spaceAfter=4
    )
    
    contact_style = ParagraphStyle(
        'ResumeContact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=11,
        textColor=colors.HexColor('#4b5563'),
        alignment=1, # Center
        spaceAfter=12
    )
    
    section_heading_style = ParagraphStyle(
        'ResumeSectionHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=13,
        textColor=colors.HexColor('#111827'),
        spaceBefore=8,
        spaceAfter=3,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'ResumeBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=9,
        leading=12.5,
        textColor=colors.HexColor('#374151'),
        spaceAfter=6
    )
    
    bullet_style = ParagraphStyle(
        'ResumeBullet',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=9,
        leading=12.5,
        textColor=colors.HexColor('#374151'),
        leftIndent=12,
        firstLineIndent=-8,
        spaceAfter=3
    )
    
    meta_style_left = ParagraphStyle(
        'ResumeMetaLeft',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9.5,
        leading=12,
        textColor=colors.HexColor('#111827'),
        spaceAfter=1
    )
    
    meta_style_right = ParagraphStyle(
        'ResumeMetaRight',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=12,
        textColor=colors.HexColor('#4b5563'),
        alignment=2, # Right
        spaceAfter=1
    )
    
    story = []
    
    # 1. Header (Name & Contact)
    name = resume_data.get('name', 'Candidate Profile')
    story.append(Paragraph(name, title_style))
    
    contact_info = []
    email = resume_data.get('email')
    if email: contact_info.append(email)
    phone = resume_data.get('phone')
    if phone: contact_info.append(phone)
    
    story.append(Paragraph("  |  ".join(contact_info), contact_style))
    
    # Divider helper function
    def add_section_divider(title):
        story.append(Spacer(1, 4))
        story.append(Paragraph(title.upper(), section_heading_style))
        from reportlab.platypus import Table, TableStyle
        line_table = Table([['']], colWidths=[540], rowHeights=[1])
        line_table.setStyle(TableStyle([
            ('LINEBELOW', (0,0), (-1,-1), 0.75, colors.HexColor('#e5e7eb')),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(line_table)
        story.append(Spacer(1, 4))
        
    # 2. Summary
    summary = resume_data.get('summary')
    if summary:
        add_section_divider("Professional Summary")
        story.append(Paragraph(summary, body_style))
        
    # 3. Skills
    skills = resume_data.get('skills', [])
    if skills:
        add_section_divider("Technical Skills")
        story.append(Paragraph("<b>Skills:</b> " + ", ".join(skills), body_style))
        
    # 4. Work Experience
    experience = resume_data.get('experience', [])
    if experience:
        add_section_divider("Professional Experience")
        from reportlab.platypus import Table, TableStyle
        for exp in experience:
            exp_elements = []
            company = exp.get('company', '')
            role = exp.get('role', '')
            duration = exp.get('duration', '')
            
            title_text = f"<b>{company}</b>" if company else ""
            if role:
                title_text += f" &ndash; <i>{role}</i>"
                
            meta_table = Table(
                [[Paragraph(title_text, meta_style_left), Paragraph(duration, meta_style_right)]],
                colWidths=[380, 160]
            )
            meta_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
                ('TOPPADDING', (0,0), (-1,-1), 2),
            ]))
            exp_elements.append(meta_table)
            
            for h in exp.get('highlights', []):
                exp_elements.append(Paragraph(f"&bull; {h}", bullet_style))
                
            exp_elements.append(Spacer(1, 4))
            story.append(KeepTogether(exp_elements))
            
    # 5. Projects
    projects = resume_data.get('projects', [])
    if projects:
        add_section_divider("Projects")
        from reportlab.platypus import Table, TableStyle
        for proj in projects:
            proj_elements = []
            title = proj.get('title', '')
            desc = proj.get('description', '')
            
            title_text = f"<b>{title}</b>"
            if desc:
                title_text += f" &ndash; <i>{desc}</i>"
                
            meta_table = Table(
                [[Paragraph(title_text, meta_style_left), Paragraph("", meta_style_right)]],
                colWidths=[400, 140]
            )
            meta_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
                ('TOPPADDING', (0,0), (-1,-1), 2),
            ]))
            proj_elements.append(meta_table)
            
            for h in proj.get('highlights', []):
                proj_elements.append(Paragraph(f"&bull; {h}", bullet_style))
                
            proj_elements.append(Spacer(1, 4))
            story.append(KeepTogether(proj_elements))
            
    # 6. Education
    education = resume_data.get('education', [])
    if education:
        add_section_divider("Education")
        from reportlab.platypus import Table, TableStyle
        for edu in education:
            edu_elements = []
            inst = edu.get('institution', '')
            deg = edu.get('degree', '')
            maj = edu.get('major', '')
            year = edu.get('graduation_year', '')
            
            degree_str = f"{deg} in {maj}" if maj else deg
            text_left = f"<b>{inst}</b>"
            if degree_str:
                text_left += f" &ndash; {degree_str}"
                
            meta_table = Table(
                [[Paragraph(text_left, meta_style_left), Paragraph(year, meta_style_right)]],
                colWidths=[420, 120]
            )
            meta_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
                ('TOPPADDING', (0,0), (-1,-1), 2),
            ]))
            edu_elements.append(meta_table)
            edu_elements.append(Spacer(1, 3))
            story.append(KeepTogether(edu_elements))
            
    doc.build(story)
    buffer.seek(0)
    return buffer


def generate_resume_docx(resume_data: dict) -> io.BytesIO:
    doc = docx.Document()
    
    # 0.5 inch margins for standard resume layout
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    # Typography configuration helper
    def set_font(run, name="Arial", size_pt=9.5, bold=False, italic=False, color_rgb="333333"):
        run.font.name = name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = docx.shared.RGBColor.from_string(color_rgb)
        
    # Title / Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(resume_data.get('name', 'Candidate Profile'))
    set_font(run_name, size_pt=18, bold=True, color_rgb="000000")
    p_name.paragraph_format.space_after = Pt(2)
    
    # Contact details
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_parts = []
    if resume_data.get('email'): contact_parts.append(resume_data['email'])
    if resume_data.get('phone'): contact_parts.append(resume_data['phone'])
    run_contact = p_contact.add_run("  |  ".join(contact_parts))
    set_font(run_contact, size_pt=9, color_rgb="555555")
    p_contact.paragraph_format.space_after = Pt(12)
    
    # Divider and heading builder
    def add_section_header(title):
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(8)
        p_hdr.paragraph_format.space_after = Pt(2)
        p_hdr.paragraph_format.keep_with_next = True
        run_hdr = p_hdr.add_run(title.upper())
        set_font(run_hdr, size_pt=10, bold=True, color_rgb="111111")
        
        # Add thin border bottom under paragraph
        pBrd = parse_xml(f'<w:pBrd {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="E5E7EB"/></w:pBrd>')
        p_hdr._p.get_or_add_pPr().append(pBrd)
        
    # Summary
    summary = resume_data.get('summary')
    if summary:
        add_section_header("Professional Summary")
        p_sum = doc.add_paragraph()
        p_sum.paragraph_format.space_after = Pt(6)
        p_sum.paragraph_format.line_spacing = 1.15
        run_sum = p_sum.add_run(summary)
        set_font(run_sum, size_pt=9.5)
        
    # Skills
    skills = resume_data.get('skills', [])
    if skills:
        add_section_header("Technical Skills")
        p_skills = doc.add_paragraph()
        p_skills.paragraph_format.space_after = Pt(6)
        run_bold = p_skills.add_run("Skills: ")
        set_font(run_bold, size_pt=9.5, bold=True)
        run_text = p_skills.add_run(", ".join(skills))
        set_font(run_text, size_pt=9.5)
        
    # Experience
    experience = resume_data.get('experience', [])
    if experience:
        add_section_header("Professional Experience")
        for exp in experience:
            company = exp.get('company', '')
            role = exp.get('role', '')
            duration = exp.get('duration', '')
            
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(5.5)
            table.columns[1].width = Inches(2.0)
            
            cell_left = table.cell(0, 0)
            p_left = cell_left.paragraphs[0]
            p_left.paragraph_format.space_after = Pt(2)
            p_left.paragraph_format.keep_with_next = True
            run_comp = p_left.add_run(company)
            set_font(run_comp, size_pt=9.5, bold=True, color_rgb="000000")
            if role:
                run_sep = p_left.add_run(" - ")
                set_font(run_sep, size_pt=9.5, color_rgb="555555")
                run_role = p_left.add_run(role)
                set_font(run_role, size_pt=9.5, italic=True, color_rgb="333333")
                
            cell_right = table.cell(0, 1)
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.space_after = Pt(2)
            run_dur = p_right.add_run(duration)
            set_font(run_dur, size_pt=9.5, color_rgb="555555")
            
            tblPr = table._tbl.tblPr
            tblBorders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>')
            tblPr.append(tblBorders)
            
            for h in exp.get('highlights', []):
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_before = Pt(0)
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.paragraph_format.left_indent = Inches(0.25)
                run_bullet = p_bullet.add_run(h)
                set_font(run_bullet, size_pt=9.5)
                
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_after = Pt(4)
            
    # Projects
    projects = resume_data.get('projects', [])
    if projects:
        add_section_header("Projects")
        for proj in projects:
            title = proj.get('title', '')
            desc = proj.get('description', '')
            
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(5.5)
            table.columns[1].width = Inches(2.0)
            
            cell_left = table.cell(0, 0)
            p_left = cell_left.paragraphs[0]
            p_left.paragraph_format.space_after = Pt(2)
            p_left.paragraph_format.keep_with_next = True
            run_title = p_left.add_run(title)
            set_font(run_title, size_pt=9.5, bold=True, color_rgb="000000")
            if desc:
                run_sep = p_left.add_run(" - ")
                set_font(run_sep, size_pt=9.5)
                run_desc = p_left.add_run(desc)
                set_font(run_desc, size_pt=9.5, italic=True)
                
            cell_right = table.cell(0, 1)
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.space_after = Pt(2)
            
            tblPr = table._tbl.tblPr
            tblBorders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>')
            tblPr.append(tblBorders)
            
            for h in proj.get('highlights', []):
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.paragraph_format.left_indent = Inches(0.25)
                run_bullet = p_bullet.add_run(h)
                set_font(run_bullet, size_pt=9.5)
                
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_after = Pt(4)
            
    # Education
    education = resume_data.get('education', [])
    if education:
        add_section_header("Education")
        for edu in education:
            inst = edu.get('institution', '')
            deg = edu.get('degree', '')
            maj = edu.get('major', '')
            year = edu.get('graduation_year', '')
            
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(5.5)
            table.columns[1].width = Inches(2.0)
            
            cell_left = table.cell(0, 0)
            p_left = cell_left.paragraphs[0]
            p_left.paragraph_format.space_after = Pt(2)
            run_inst = p_left.add_run(inst)
            set_font(run_inst, size_pt=9.5, bold=True, color_rgb="000000")
            
            degree_str = f"{deg} in {maj}" if maj else deg
            if degree_str:
                run_sep = p_left.add_run(" - ")
                set_font(run_sep, size_pt=9.5)
                run_deg = p_left.add_run(degree_str)
                set_font(run_deg, size_pt=9.5)
                
            cell_right = table.cell(0, 1)
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.space_after = Pt(2)
            run_year = p_right.add_run(year)
            set_font(run_year, size_pt=9.5, color_rgb="555555")
            
            tblPr = table._tbl.tblPr
            tblBorders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>')
            tblPr.append(tblBorders)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
